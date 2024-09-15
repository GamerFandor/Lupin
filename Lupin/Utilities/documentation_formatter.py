# Imports
from docx import Document



# Create heading
def heading(document: Document, title : str) -> None:
    document.add_heading(title, level=1)



# Create subheading
def subheading(document: Document, title : str) -> None:
    document.add_heading(title, level=2)



# Create paragraph
def paragraph(document: Document, text : str) -> None:
    document.add_paragraph(text)



# Create unordered list
def unordered_list(document: Document, items : list) -> None:
    for item in items:
        document.add_paragraph(item, style='ListBullet')



# Create ordered list
def ordered_list(document: Document, items : list) -> None:
    for item in items:
        document.add_paragraph(item, style='ListNumber')



# Create table
def table(document: Document, header_row_data : list, data : list) -> None:
    table = document.add_table(rows=1, cols=len(header_row_data))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(header_row_data):
        hdr_cells[i].text = header
    for row in data:
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = cell